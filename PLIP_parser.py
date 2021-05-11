#### A PROGRAM TO PARSE THE PLIP RESULT FILES 
#### FOR PROTEIN-PROTEIN and PROTEIN-LIGAND INTERACTIONS 
#### AND MAKING A TABLE IN MS WORD FILE

# step1: converting triple letters with single letter AA
import os
import docx

print('>>> THIS PROGRAM PARSES THE TXT FILES PRODUCED FROM PLIP WEBSERVER.\nIT WORKS WITH \n1- HYDROPHOBIC INTERACTIONS\n2- H-BONDS\n3- SALT BRIDGES\n4- Pi-stacking\n(until now).\nIF YOU HAVE OTHER INTERACTIONS IN THE RESULT FILE AND THE PROGRAM DID NOT DETECT THEM, \nCONTACT ME AT MY GITHUB REPO: https://github.com/hima111997/Plip_parser \n\n\n')

pod = int(input('>>> Type 0 for protein-protein interaction \nor 1 for protein-ligand interaction: \n'))
### 0: protein - 1: drug
aa_names = {'ALA':'A', 'ARG':'R', 'ASN':'N', 'ASP':'D', 'CYS':'C', 'GLN':'Q',
           'GLU':'E', 'GLY':'G', 'HIS':'H', 'ILE':'I', 'LEU':'L', 'LYS':'K',
           'MET':'M', 'PHE':'F', 'PRO':'P', 'SER':'S', 'THR':'T', 'TRP':'W',
           'TYR':'Y', 'VAL':'V'}
#dir_ = r'D:\mds\protein-ligand'
dir_ = input('\n\n>>> Type the destination that \ncontains the folders having the result files: ' )
table_name = input('\n\n>>> Type a name for the output MS table: ' )
folders = [(dir_ + '\\' + f , os.listdir(dir_ + '\\' + f)) for f in os.listdir(dir_) if os.path.isdir(dir_ + '\\' + f)]


files = [(d, [f for f in fs if f.endswith('.txt')]) for d, fs in folders]
new_files = []
for d, i in files:
    for txts in i:
        with open(d+'\\'+txts) as f:
            data = f.read()
            for aa in aa_names:
                data = data.replace(aa, aa_names[aa])    
            with open(d+'\\'+txts+'_mod.txt','w') as f2:
                f2.write(data)
            new_files.append(d+'\\'+txts+'_mod.txt')
            
            
# step2: Extracting interacting AA
def parsing(table, pod, sb = 0):
    AA = []
    lines = table.splitlines()[3::2]
    #print(lines)
    for l in lines:        
        elements = l.split('|')
        
        if pod == 0:
            if sb == 0:
                AA.append((elements[1].strip(), elements[2].strip(), elements[4].strip(), elements[5].strip()))
            else:
                #print('Salt Bridges found')
                AA.append((elements[1].strip(), elements[2].strip(), elements[5].strip(), elements[6].strip()))
        else:
            AA.append((elements[1].strip(), elements[2].strip()))
        ####### AA number (R), AA letter (R), AA number (L), AA letter (L)
    return AA

type_AA = {}
file_type_AA = {}
for txt in new_files:
    with open(txt) as f:
        data = f.read()
    num_interactions = int(data.count('**')/2)
    interactions = []
    if num_interactions > 1:
        start = 0
        for inter in range(num_interactions):
            type_interactions = data.find('**', start)        
            interactions.append((data.find('**', type_interactions+1)+2 , data[type_interactions+2 : data.find('**', type_interactions+1)]))
            start = data.find('**', type_interactions+1)+2

        for idx, interaction in enumerate(interactions[:-1]):
            sb=0
            if 'Salt' in interaction[1]:
                sb = 1
            table = data[interaction[0]:data.find(interactions[idx+1][1])-2]
            #print(table.strip())
            type_AA[interaction[1]] = parsing(table.strip(), sb=sb, pod=pod)
        sb=0
        if 'Salt' in interactions[-1][1]:
            sb = 1
        table = data[interactions[-1][0]:]
        type_AA[interactions[-1][1]] = parsing(table.strip(), sb=sb, pod=pod)
    else:
        type_interactions = data.find('**')
        interactions.append((data.find('**', type_interactions+1)+2 , data[type_interactions+2 : data.find('**', type_interactions+1)]))
        sb=0
        if 'Salt' in interaction[1]:
            sb = 1
        table = data[interactions[0][0]:]
        type_AA[interactions[0][1]] = parsing(table.strip(), sb=sb, pod=pod)
        
    file_type_AA[txt.split('\\')[-2]] = type_AA
    type_AA = {}
#print(file_type_AA)


#step3: create a MS word file containing the table


  
rows = len(file_type_AA)
#columns = max([len(i) for i in file_type_AA.values()])
set_types = set()
for i in file_type_AA.values():
    set_types.update(i.keys())
set_types = list(set_types)


# Create an instance of a word document
doc = docx.Document()

if pod == 0:
    MS_table = doc.add_table(rows=rows+1, cols=len(set_types) * 3 + 1)
    row = MS_table.rows[0].cells
    row[0].text = 'Cluster Number'
    # protein-protein
    for i in range(len(set_types)):
        row[1 + (3*i)].text = 'Number of {}'.format(set_types[i])
        row[2 + (3*i)].text = 'Amino Acids in receptor'
        row[3 + (3*i)].text = 'Amino Acids in ligand'

    # Adding data from the list to the table
    for idx, (k, v) in enumerate(file_type_AA.items()):

        # Adding a row and then adding data in it.
        #row = table.add_row().cells
        row = MS_table.rows[idx+1].cells
        # Converting id to string as table can only take string input
        row[0].text = k
        for i in range(len(set_types)):        
            number = v.get(set_types[i], '0')
            if number == '0':
                row[1 + (3*i)].text = '0'
                row[2 + (3*i)].text= 'None'
                row[3 + (3*i)].text = 'None'
            else:
                row[1 + (3*i)].text = str(len(number))
                row[2 + (3*i)].text = ' - '.join([i[1]+i[0] for i in number])
                row[3 + (3*i)].text = ' - '.join([i[3]+i[2] for i in number])
else:
    MS_table = doc.add_table(rows=rows+1, cols=len(set_types) * 2 + 1)
    row = MS_table.rows[0].cells
    row[0].text = 'Cluster Number'
    # protein-drug    
    for i in range(len(set_types)):
        row[1 + (2*i)].text = 'Number of {}'.format(set_types[i])
        row[2 + (2*i)].text = 'Amino Acids in receptor'        

    # Adding data from the list to the table
    for idx, (k, v) in enumerate(file_type_AA.items()):

        # Adding a row and then adding data in it.
        #row = table.add_row().cells
        row = MS_table.rows[idx+1].cells
        # Converting id to string as table can only take string input
        row[0].text = k
        for i in range(len(set_types)):        
            number = v.get(set_types[i], '0')
            if number == '0':
                row[1 + (2*i)].text = '0'
                row[2 + (2*i)].text= 'None'                
            else:
                row[1 + (2*i)].text = str(len(number))
                row[2 + (2*i)].text = ' - '.join([i[1]+i[0] for i in number])                
 # Now save the document to a location
doc.save(dir_ + '\\' + '{}.docx'.format(table_name))
print('\n\n>>> You can see the output file at: {}'.format(dir_))
print('\n\n>>> Thank You for using This Program.\n if You have any recommendations or found a problem,\n contact me at my GITHUB repo: \nhttps://github.com/hima111997/Plip_parser \n')
input('\n\n>>> Press any key to close..')
